"""Guards for figure and table fidelity in the Chapter 4 DOCX pipeline.

Three separate silent-data-loss bugs motivated these tests:

1. build_ch4_docx.py had no image branch, so `![alt](path)` was emitted as
   literal body text and no figure ever reached the DOCX.
2. build_ch4_docx.py treated any line containing '|' as a table row, which
   swallowed prose with inline math such as $P(d|p)$ or $|matched|/|all|$.
3. inject_ch4_into_main_thesis.py rebuilt the document from a blank
   `Document()` and copied only `run.text`, discarding embedded images,
   tables outside Chapter 4, and section properties.

Each failure produced a valid-looking DOCX with content missing, so these
tests assert on structure rather than on the build merely succeeding.
"""

import copy
import io
import zipfile

import docx
import pytest
from PIL import Image
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree

from scripts.utils.inject_ch4_into_main_thesis import (
    _find_chapter_bounds,
    _remap_images,
    replace_ch4_in_main_doc,
)
from scripts.build_ch4_docx import add_figure


A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def _png_bytes():
    """A small real PNG. python-docx parses the header for dimensions, so a
    hand-written byte string is not good enough."""
    buffer = io.BytesIO()
    Image.new('RGB', (4, 4), (200, 60, 60)).save(buffer, format='PNG')
    return buffer.getvalue()


@pytest.fixture
def png(tmp_path):
    path = tmp_path / 'fig.png'
    path.write_bytes(_png_bytes())
    return path


def _blips(document):
    return list(document.element.body.iter(f'{{{A_NS}}}blip'))


# ── build_ch4_docx: figures ────────────────────────────────────────────────

def test_add_figure_embeds_a_real_picture(tmp_path, png, monkeypatch):
    import scripts.build_ch4_docx as builder
    monkeypatch.setattr(builder, 'MD_PATH', tmp_path / 'chapter.md')

    doc = docx.Document()
    add_figure(doc, 'รูปที่ 4.A ทดสอบ', 'fig.png')

    assert len(_blips(doc)) == 1
    media = [p for p in doc.part.package.parts if 'media' in str(p.partname)]
    assert len(media) == 1


def test_add_figure_writes_a_caption(tmp_path, png, monkeypatch):
    import scripts.build_ch4_docx as builder
    monkeypatch.setattr(builder, 'MD_PATH', tmp_path / 'chapter.md')

    doc = docx.Document()
    add_figure(doc, 'รูปที่ 4.D การแยกกันของคะแนนและอันดับ', 'fig.png')

    text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'รูปที่ 4.D การแยกกันของคะแนนและอันดับ' in text


def test_add_figure_raises_on_missing_file(tmp_path, monkeypatch):
    """Skipping a missing figure would leave the chapter referencing nothing."""
    import scripts.build_ch4_docx as builder
    monkeypatch.setattr(builder, 'MD_PATH', tmp_path / 'chapter.md')

    doc = docx.Document()
    with pytest.raises(FileNotFoundError, match='figure not found'):
        add_figure(doc, 'missing', 'no_such_figure.png')


def test_table_detection_requires_a_leading_pipe():
    """Regression guard: prose with inline math must not be eaten as a table."""
    import inspect

    source = inspect.getsource(docx and __import__(
        'scripts.build_ch4_docx', fromlist=['convert_md_to_docx']
    ).convert_md_to_docx)
    assert "line.strip().startswith('|')" in source
    assert "if '|' in line and not in_code_block" not in source


@pytest.mark.parametrize('prose', [
    'การทดลองนี้เปรียบเทียบวิธีคำนวณ $P(d|p)$ สามแบบ',
    '- **Keyword (Graded)** — สัดส่วนคีย์เวิร์ด ($|matched|/|all|$)',
    "*Cohen's $d$ อยู่ในระดับ negligible ($|d| \\le 0.196$)*",
])
def test_inline_math_with_pipes_survives_conversion(tmp_path, prose, monkeypatch):
    import scripts.build_ch4_docx as builder

    md = tmp_path / 'chapter.md'
    md.write_text(f'## หัวข้อ\n\n{prose}\n', encoding='utf-8')
    out = tmp_path / 'out.docx'
    monkeypatch.setattr(builder, 'MD_PATH', md)
    builder.convert_md_to_docx(md, out)

    text = '\n'.join(p.text for p in docx.Document(out).paragraphs)
    # Bullet markers and bold markup are restyled, so compare on a distinctive
    # fragment rather than the raw line.
    needle = prose.split('—')[-1].strip().lstrip('*').strip()
    assert needle[:20] in text


# ── injection: images, tables, and untouched regions ──────────────────────

def _make_main_doc(path, ch4_body=('เนื้อหาบทที่ 4 เดิม',)):
    doc = docx.Document()
    doc.add_paragraph('หน้าปกและบทที่ 1-3')
    doc.add_paragraph('บทที่ 4 ผลการวิจัย')
    for line in ch4_body:
        doc.add_paragraph(line)
    doc.add_paragraph('บทที่ 5 สรุปผล')
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'SUMMARY OF CASE-BASED EMPIRICAL EVIDENCE'
    doc.add_paragraph('ท้ายบทที่ 5')
    doc.save(path)
    return path


def _make_ch4_fragment(path, png_path):
    doc = docx.Document()
    doc.add_paragraph('บทที่ 4 ผลการวิจัย (ฉบับใหม่)')
    doc.add_paragraph('4.5.4 ผลการเปรียบเทียบ Soft กับ Hard')
    doc.add_paragraph().add_run().add_picture(str(png_path), width=Inches(2))
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'โหมด'
    table.cell(1, 0).text = 'Keyword (Graded)'
    doc.save(path)
    return path


def test_injection_preserves_embedded_images(tmp_path, png):
    main = _make_main_doc(tmp_path / 'main.docx')
    ch4 = _make_ch4_fragment(tmp_path / 'ch4.docx', png)
    out = tmp_path / 'out.docx'

    replace_ch4_in_main_doc(str(main), str(out), ch4_path=str(ch4))

    result = docx.Document(out)
    assert len(_blips(result)) == 1
    media = [p for p in result.part.package.parts if 'media' in str(p.partname)]
    assert len(media) == 1


def test_injected_image_relationship_ids_resolve(tmp_path, png):
    """A deep-copied run keeps the source rId; unremapped, Word calls it corrupt."""
    main = _make_main_doc(tmp_path / 'main.docx')
    ch4 = _make_ch4_fragment(tmp_path / 'ch4.docx', png)
    out = tmp_path / 'out.docx'

    replace_ch4_in_main_doc(str(main), str(out), ch4_path=str(ch4))

    with zipfile.ZipFile(out) as archive:
        assert archive.testzip() is None
        document = etree.fromstring(archive.read('word/document.xml'))
        rels = etree.fromstring(archive.read('word/_rels/document.xml.rels'))

    declared = {rel.get('Id') for rel in rels}
    used = {
        blip.get(f'{{{R_NS}}}embed')
        for blip in document.iter(f'{{{A_NS}}}blip')
    }
    used.discard(None)

    assert used, 'no image references found in the injected document'
    assert used <= declared, f'dangling relationship ids: {used - declared}'


def test_injection_keeps_tables_outside_chapter_four(tmp_path, png):
    main = _make_main_doc(tmp_path / 'main.docx')
    ch4 = _make_ch4_fragment(tmp_path / 'ch4.docx', png)
    out = tmp_path / 'out.docx'

    replace_ch4_in_main_doc(str(main), str(out), ch4_path=str(ch4))

    result = docx.Document(out)
    cells = [
        cell.text
        for table in result.tables
        for row in table.rows
        for cell in row.cells
    ]
    assert any('SUMMARY OF CASE-BASED' in text for text in cells)


def test_injection_leaves_surrounding_chapters_untouched(tmp_path, png):
    main = _make_main_doc(tmp_path / 'main.docx')
    ch4 = _make_ch4_fragment(tmp_path / 'ch4.docx', png)
    out = tmp_path / 'out.docx'

    replace_ch4_in_main_doc(str(main), str(out), ch4_path=str(ch4))

    texts = [p.text for p in docx.Document(out).paragraphs]
    assert texts[0] == 'หน้าปกและบทที่ 1-3'
    assert 'ท้ายบทที่ 5' in texts
    assert 'เนื้อหาบทที่ 4 เดิม' not in texts, 'old Chapter 4 was not removed'
    assert '4.5.4 ผลการเปรียบเทียบ Soft กับ Hard' in texts


def test_injection_does_not_import_the_fragment_section_properties(tmp_path, png):
    """Section setup belongs to the main document, not the Chapter 4 fragment."""
    main = _make_main_doc(tmp_path / 'main.docx')
    ch4 = _make_ch4_fragment(tmp_path / 'ch4.docx', png)
    out = tmp_path / 'out.docx'

    replace_ch4_in_main_doc(str(main), str(out), ch4_path=str(ch4))

    result = docx.Document(out)
    assert len(result.sections) == len(docx.Document(main).sections)
    sect_prs = result.element.body.findall(qn('w:sectPr'))
    assert len(sect_prs) == 1


def test_injection_raises_when_chapter_heading_is_missing(tmp_path, png):
    doc = docx.Document()
    doc.add_paragraph('ไม่มีหัวบท')
    main = tmp_path / 'main.docx'
    doc.save(main)
    ch4 = _make_ch4_fragment(tmp_path / 'ch4.docx', png)

    with pytest.raises(ValueError, match='could not locate'):
        replace_ch4_in_main_doc(str(main), str(tmp_path / 'out.docx'),
                                ch4_path=str(ch4))


def test_find_chapter_bounds_returns_headings_in_order(tmp_path):
    main = _make_main_doc(tmp_path / 'main.docx')
    doc = docx.Document(main)

    ch4, ch5 = _find_chapter_bounds(doc)

    assert 'บทที่ 4' in ch4.text
    assert 'บทที่ 5' in ch5.text
    body = list(doc.element.body)
    assert body.index(ch4._p) < body.index(ch5._p)


def test_remap_images_rewrites_the_relationship_id(tmp_path, png):
    source = docx.Document()
    source.add_paragraph().add_run().add_picture(str(png), width=Inches(1))
    target = docx.Document()

    element = copy.deepcopy(source.element.body[0])
    original_rid = element.find(f'.//{{{A_NS}}}blip').get(f'{{{R_NS}}}embed')

    _remap_images(element, source.part, target.part)
    new_rid = element.find(f'.//{{{A_NS}}}blip').get(f'{{{R_NS}}}embed')

    # The id may coincidentally match the source's (two blank documents hand
    # out the same next id), so assert on resolution, not on the id changing.
    assert new_rid in target.part.rels, (
        f'rId {new_rid!r} (was {original_rid!r}) does not resolve in the target'
    )
    assert target.part.rels[new_rid].target_part.blob == _png_bytes()
