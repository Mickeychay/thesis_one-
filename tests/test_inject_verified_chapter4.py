from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.inject_verified_chapter4 import marker_elements, replace_chapter


def _set_page_number_start(section_properties, start: int, fmt: str = "decimal") -> None:
    page_number = OxmlElement("w:pgNumType")
    page_number.set(qn("w:start"), str(start))
    page_number.set(qn("w:fmt"), fmt)
    section_properties.append(page_number)


def _section_break(paragraph, *, page_start: int | None = None) -> None:
    properties = paragraph._p.get_or_add_pPr()
    section = OxmlElement("w:sectPr")
    page_size = OxmlElement("w:pgSz")
    page_size.set(qn("w:w"), "11906")
    page_size.set(qn("w:h"), "16838")
    section.append(page_size)
    if page_start is not None:
        _set_page_number_start(section, page_start)
    properties.append(section)


def _section_relationship_references(document: Document) -> list[list[tuple[str, str, str]]]:
    references = []
    for section in document.sections:
        current = []
        for tag in ("w:headerReference", "w:footerReference"):
            for reference in section._sectPr.findall(qn(tag)):
                current.append((
                    tag,
                    reference.get(qn("w:type")),
                    reference.get(qn("r:id")),
                ))
        references.append(current)
    return references


def test_replaces_chapter_tables_and_preserves_section_break(tmp_path: Path):
    source_path = tmp_path / "source.docx"
    chapter_path = tmp_path / "chapter.docx"
    output_path = tmp_path / "output.docx"

    source = Document()
    source.add_paragraph("บทที่ 3")
    source.add_paragraph("เนื้อหาบทที่สาม")
    source.add_paragraph("บทที่ 4")
    source.add_paragraph("เนื้อหาเก่า")
    source.add_table(rows=1, cols=1).cell(0, 0).text = "ตารางเก่า"
    section_paragraph = source.add_paragraph()
    _section_break(section_paragraph, page_start=21)
    source.add_paragraph("บทที่ 5")
    source.add_paragraph("เนื้อหาบทที่ห้า")
    final_section = source.sections[-1]
    _set_page_number_start(final_section._sectPr, 25)
    final_section.header.is_linked_to_previous = False
    final_section.footer.is_linked_to_previous = False
    final_section.header.paragraphs[0].text = "HEADER FIVE"
    final_section.footer.paragraphs[0].text = "FOOTER FIVE"
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "false")
    source.settings._element.append(update_fields)
    source.save(source_path)

    chapter = Document()
    chapter.add_paragraph("บทที่ 4")
    chapter.add_paragraph("ผลชุดใหม่ 220 เคส")
    chapter.add_table(rows=1, cols=2).rows[0].cells[0].text = "ตารางใหม่"
    chapter.save(chapter_path)

    before = Document(source_path)
    before_sections = len(before.sections)
    before_relationships = _section_relationship_references(before)
    result = replace_chapter(source_path, chapter_path, output_path)
    after = Document(output_path)

    assert len(marker_elements(after, "บทที่ 4")) == 1
    assert len(marker_elements(after, "บทที่ 5")) == 1
    assert "ผลชุดใหม่ 220 เคส" in [paragraph.text for paragraph in after.paragraphs]
    assert "เนื้อหาเก่า" not in [paragraph.text for paragraph in after.paragraphs]
    assert [cell.text for table in after.tables for row in table.rows for cell in row.cells] == [
        "ตารางใหม่",
        "",
    ]
    assert len(after.sections) == before_sections
    assert result["preserved_section_breaks"] == 1
    assert result["cleared_page_number_starts"] == 2

    chapter_start = marker_elements(after, "บทที่ 4")[0]
    properties = chapter_start.get_or_add_pPr()
    assert properties.find(qn("w:pageBreakBefore")).get(qn("w:val")) == "true"
    assert properties.find(qn("w:pStyle")).get(qn("w:val")) == "Heading1"
    assert properties.find(qn("w:outlineLvl")).get(qn("w:val")) == "0"

    page_number_elements = list(after.element.body.iter(qn("w:pgNumType")))
    assert len(page_number_elements) == 2
    assert all(element.get(qn("w:start")) is None for element in page_number_elements)
    assert all(element.get(qn("w:fmt")) == "decimal" for element in page_number_elements)
    assert _section_relationship_references(after) == before_relationships
    assert after.sections[-1].header.paragraphs[0].text == "HEADER FIVE"
    assert after.sections[-1].footer.paragraphs[0].text == "FOOTER FIVE"

    preserved_update_fields = after.settings._element.find(qn("w:updateFields"))
    assert preserved_update_fields is not None
    assert preserved_update_fields.get(qn("w:val")) == "false"


def test_does_not_add_global_field_update_setting(tmp_path: Path):
    source_path = tmp_path / "source.docx"
    chapter_path = tmp_path / "chapter.docx"
    output_path = tmp_path / "output.docx"

    source = Document()
    source.add_paragraph("บทที่ 4")
    source.add_paragraph("เนื้อหาเก่า")
    source.add_paragraph("บทที่ 5")
    source.save(source_path)

    chapter = Document()
    chapter.add_paragraph("บทที่ 4")
    chapter.add_paragraph("เนื้อหาใหม่")
    chapter.save(chapter_path)

    replace_chapter(source_path, chapter_path, output_path)
    after = Document(output_path)
    assert after.settings._element.find(qn("w:updateFields")) is None


def test_rejects_chapter_that_contains_chapter_five(tmp_path: Path):
    source_path = tmp_path / "source.docx"
    chapter_path = tmp_path / "chapter.docx"
    output_path = tmp_path / "output.docx"
    source = Document()
    source.add_paragraph("บทที่ 4")
    source.add_paragraph("บทที่ 5")
    source.save(source_path)
    chapter = Document()
    chapter.add_paragraph("บทที่ 4")
    chapter.add_paragraph("บทที่ 5")
    chapter.save(chapter_path)

    try:
        replace_chapter(source_path, chapter_path, output_path)
    except ValueError as exc:
        assert "must not contain" in str(exc)
    else:
        raise AssertionError("Expected chapter validation to fail")
