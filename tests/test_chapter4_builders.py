import json
from pathlib import Path

import pandas as pd
import pytest
from docx import Document
from docx.oxml.ns import qn

import scripts.build_chapter4_evidence as evidence_builder
from scripts.build_chapter4_evidence import (
    EvidencePaths,
    EXPECTED_MODELS,
    artifact_date_statement,
    artifact_provenance_window,
    evaluation_protocol_rows,
    l2_slice_summary,
    sensitivity_table_rows,
    sha256,
    validate_ground_truth,
    validate_ground_truth_audit,
)
from scripts.build_verified_chapter4 import (
    FONT_NAME,
    add_heading,
    parse_table,
    split_markdown_row,
    verify_source_manifest,
)
from scripts.inject_verified_chapter4 import DEFAULT_CHAPTER


def make_case(case_id: str, split: str, evaluation_slice: str | None = None) -> dict:
    case = {
        "case_id": case_id,
        "split": split,
        "case_description": case_id,
        "expected_diagnosis": {"problem_list": [{"code": "1001"}]},
    }
    if evaluation_slice:
        case["evaluation_slice"] = evaluation_slice
    return case


def complete_ground_truth() -> dict:
    cases = [make_case(f"TRAIN_{index:03d}", "train") for index in range(125)]
    cases += [make_case(f"TEST_{index:03d}", "test") for index in range(75)]
    cases += [
        make_case(f"ADV_{index:03d}", "test", "adversarial_test")
        for index in range(20)
    ]
    for case in cases[70:106]:
        case["augmentation"] = {"type": "paraphrase"}
    for case in cases[106:115]:
        case["augmentation"] = {"type": "complexity_escalation"}
    for case in cases[115:125]:
        case["augmentation"] = {"type": "complexity_reduction"}
    for case in cases[155:163]:
        case["augmentation"] = {"type": "paraphrase"}
    cases[163]["augmentation"] = {"type": "complexity_escalation"}
    for case in cases[164:200]:
        case["augmentation"] = {"type": "polarity"}
    for case in cases[200:220]:
        case["augmentation"] = {"type": "adversarial"}
    return {
        "metadata": {
            "total_cases": 220,
            "train_cases": 125,
            "test_cases": 95,
            "adversarial_test_cases": 20,
            "split_method": "family_level_stratified_by_category",
            "split_seed": 42,
            "near_duplicate_resolution_threshold": 0.90,
        },
        "cases": cases,
    }


def test_ground_truth_validator_accepts_final_protocol():
    result = validate_ground_truth(complete_ground_truth())
    assert result["split_counts"] == {"train": 125, "test": 95}
    assert result["slice_counts"] == {
        "standard_test": 75,
        "adversarial_test": 20,
    }


def test_ground_truth_validator_rejects_duplicate_ids():
    payload = complete_ground_truth()
    payload["cases"][-1]["case_id"] = payload["cases"][-2]["case_id"]
    with pytest.raises(ValueError, match="duplicate case IDs"):
        validate_ground_truth(payload)


@pytest.mark.parametrize(
    ("field", "value", "message"),
    [
        ("split_method", "case_level_random", "split_method"),
        ("split_seed", 7, "split_seed"),
        ("near_duplicate_resolution_threshold", 0.85, "near-duplicate threshold"),
    ],
)
def test_ground_truth_validator_rejects_wrong_split_protocol(field, value, message):
    payload = complete_ground_truth()
    payload["metadata"][field] = value
    with pytest.raises(ValueError, match=message):
        validate_ground_truth(payload)


def test_ground_truth_audit_distinguishes_original_from_generated_cases():
    ground_truth = validate_ground_truth(complete_ground_truth())
    audit = {
        "total_cases": 220,
        "split_counts": {"train": 125, "test": 95},
        "augmentation_counts": {
            "original": 100,
            "paraphrase": 44,
            "complexity_escalation": 10,
            "complexity_reduction": 10,
            "adversarial": 20,
            "polarity": 36,
        },
        "split_by_augmentation": {
            "original": {"train": 70, "test": 30},
            "paraphrase": {"train": 36, "test": 8},
            "complexity_escalation": {"train": 9, "test": 1},
            "complexity_reduction": {"train": 10},
            "adversarial": {"test": 20},
            "polarity": {"test": 36},
        },
        "missing_split": [],
        "cross_split_families": [],
        "exact_duplicates": [],
        "near_duplicates": [],
        "duplicate_threshold": 0.90,
    }
    result = validate_ground_truth_audit(audit, ground_truth)
    assert result["original_non_augmented_cases"] == 100
    assert result["generated_modified_cases"] == 120
    assert result["family_leakage_count"] == 0
    rows = evaluation_protocol_rows(result)
    assert rows[2] == [
        "เคสอ้างอิงที่ไม่ติดป้าย augmentation (92 เคสหลัก + Short 5 + Tiny 3)",
        "100 กรณี",
    ]
    assert rows[3] == [
        "เคสสร้าง/ดัดแปลง (Paraphrase 44 + Escalation 10 + Reduction 10 + Adversarial 20 + Polarity 36)",
        "120 กรณี",
    ]


def test_artifact_provenance_window_reports_cross_day_run_in_thai_calendar():
    window = artifact_provenance_window(
        {"created_at": "2026-08-08T00:30:00+07:00"},
        {"generated_at": "2026-08-08T01:15:00+07:00"},
        {"timestamp": "2026-08-07T23:51:25"},
        {"generated_at": "2026-08-08T01:20:00+07:00"},
        {"completed_at": "2026-08-08T01:05:00"},
        {"completed_at": "2026-08-08T00:04:16+07:00"},
    )
    assert window["timezone"] == "Asia/Bangkok"
    assert window["started_date_th"] == "7 สิงหาคม 2569"
    assert window["completed_date_th"] == "8 สิงหาคม 2569"
    assert len(window["events"]) == 6
    assert artifact_date_statement(window) == (
        "การรันและสร้าง artifact เริ่มเมื่อวันที่ 7 สิงหาคม 2569 "
        "และสรุปผลครบเมื่อวันที่ 8 สิงหาคม 2569"
    )


def test_artifact_provenance_window_converts_utc_before_choosing_dates():
    window = artifact_provenance_window(
        {"created_at": "2026-08-07T17:30:00Z"},
        {"generated_at": "2026-08-08T00:45:00+07:00"},
        {"timestamp": "2026-08-08T00:31:00+07:00"},
        {"generated_at": "2026-08-08T00:42:00+07:00"},
        {"completed_at": "2026-08-08T00:40:00+07:00"},
        {"completed_at": "2026-08-08T00:35:00+07:00"},
    )
    assert window["started_date_th"] == "8 สิงหาคม 2569"
    assert window["completed_date_th"] == "8 สิงหาคม 2569"
    assert artifact_date_statement(window) == (
        "artifact ทั้งหมดสร้างและสรุปผลเมื่อวันที่ 8 สิงหาคม 2569"
    )


def test_sensitivity_table_does_not_call_unexercised_zero_delta_stable():
    sensitivity = {
        "metadata": {
            "scoring_assumptions": {
                "not_exercised_parameters": {
                    "MARGIN_M": "no embeddings supplied",
                    "L1_WEIGHT_BETA": "detected problems held fixed",
                }
            }
        },
        "frame": pd.DataFrame([
            {"parameter": "MARGIN_M", "label": "m (Margin)", "value": 0.3, "is_default": True, "delta_score_pct": 0.0},
            {"parameter": "MARGIN_M", "label": "m (Margin)", "value": 0.5, "is_default": False, "delta_score_pct": 0.0},
            {"parameter": "KL_KAPPA", "label": "kappa", "value": 0.15, "is_default": True, "delta_score_pct": 0.0},
            {"parameter": "KL_KAPPA", "label": "kappa", "value": 0.30, "is_default": False, "delta_score_pct": 6.0},
        ]),
    }
    rows = sensitivity_table_rows(sensitivity)
    assert rows[0][-1] == "ไม่ถูกกระตุ้นในสมมติฐานนี้"
    assert rows[1][-1] == "ปานกลาง"


def test_evidence_build_orchestration_wires_audit_and_stress_sources(tmp_path, monkeypatch):
    ablation_dir = tmp_path / "ablation"
    sensitivity_dir = tmp_path / "sensitivity"
    ablation_dir.mkdir()
    sensitivity_dir.mkdir()
    paths = EvidencePaths(
        ground_truth=tmp_path / "ground_truth.json",
        ground_truth_audit=tmp_path / "ground_truth_audit.json",
        taxonomy=tmp_path / "taxonomy.json",
        document_metadata=tmp_path / "metadata.json",
        matrix=tmp_path / "matrix.json",
        retrieval_json=tmp_path / "retrieval.json",
        retrieval_per_case=tmp_path / "retrieval.csv",
        retrieval_significance=tmp_path / "retrieval_significance.csv",
        polarity=tmp_path / "polarity.json",
        adversarial_stress=tmp_path / "adversarial.json",
        ablation_dir=ablation_dir,
        sensitivity_dir=sensitivity_dir,
        output=tmp_path / "chapter4.md",
        manifest=tmp_path / "manifest.json",
    )
    required = [
        paths.ground_truth,
        paths.ground_truth_audit,
        paths.taxonomy,
        paths.document_metadata,
        paths.matrix,
        paths.retrieval_json,
        paths.retrieval_per_case,
        paths.retrieval_significance,
        paths.polarity,
        paths.adversarial_stress,
        paths.ablation_results,
        paths.ablation_significance,
        paths.ablation_slices,
        paths.ablation_metadata,
        paths.sensitivity_raw,
        paths.sensitivity_metadata,
    ]
    for path in required:
        path.write_text("{}", encoding="utf-8")

    ground_truth = {"cases": [], "test_cases": []}
    audit = {
        "original_non_augmented_cases": 100,
        "generated_modified_cases": 120,
        "family_leakage_count": 0,
        "exact_duplicate_count": 0,
        "cross_split_near_duplicate_count": 0,
        "near_duplicate_threshold": 0.90,
    }
    matrix = {"metadata": {"created_at": "2026-08-07T23:00:00+07:00"}, "rows_by_model": {}}
    retrieval = {"metadata": {"generated_at": "2026-08-08T01:00:00+07:00"}, "frame": pd.DataFrame()}
    polarity = {"metadata": {"timestamp": "2026-08-07T23:30:00+07:00"}}
    adversarial = {
        "metadata": {"generated_at": "2026-08-08T01:05:00+07:00"},
        "max_case_first_retrieval_difference": 0.0,
        "max_aggregate_retrieval_difference": 0.0,
    }
    ablation = {"metadata": {"completed_at": "2026-08-08T01:10:00+07:00"}}
    sensitivity = {"metadata": {"completed_at": "2026-08-08T00:05:00+07:00"}}

    monkeypatch.setattr(evidence_builder, "load_json", lambda path: {"aggregates": []} if path == paths.retrieval_json else {})
    monkeypatch.setattr(evidence_builder, "validate_ground_truth", lambda payload: ground_truth)
    monkeypatch.setattr(evidence_builder, "validate_ground_truth_audit", lambda payload, validated: audit)
    monkeypatch.setattr(evidence_builder, "validate_matrix", lambda payload, validated, selected_paths: matrix)
    monkeypatch.setattr(evidence_builder, "validate_retrieval", lambda *args: retrieval)
    monkeypatch.setattr(evidence_builder, "validate_polarity", lambda *args: polarity)
    monkeypatch.setattr(evidence_builder, "validate_adversarial_stress", lambda *args: adversarial)
    monkeypatch.setattr(evidence_builder, "validate_ablation", lambda *args: ablation)
    monkeypatch.setattr(evidence_builder, "validate_sensitivity", lambda *args: sensitivity)

    captured = {}

    def fake_markdown(
        validated_ground_truth,
        validated_audit,
        validated_matrix,
        validated_retrieval,
        validated_polarity,
        validated_ablation,
        validated_sensitivity,
        sources,
        provenance_window,
    ):
        captured["audit"] = validated_audit
        captured["sources"] = sources
        return "# บทที่ 4\n\nผลการวิจัยชุด 220 กรณี"

    monkeypatch.setattr(evidence_builder, "build_markdown", fake_markdown)
    manifest = evidence_builder.build(paths)
    assert captured["audit"] is audit
    assert "Ground-truth audit" in captured["sources"]
    assert "Adversarial stress-test summary" in captured["sources"]
    assert manifest["validation"]["original_non_augmented_cases"] == 100
    assert manifest["validation"]["adversarial_case_first_retrieval_max_abs_difference"] == 0.0
    assert paths.output.is_file()
    assert paths.manifest.is_file()


def test_markdown_table_parser_preserves_escaped_pipes():
    assert split_markdown_row(r"| a | b \| c | d |") == ["a", "b | c", "d"]
    headers, rows = parse_table(
        ["| A | B |", "|---|---|", r"| x | y \| z |"]
    )
    assert headers == ["A", "B"]
    assert rows == [["x", "y | z"]]


def test_chapter_heading_has_page_break_outline_and_direct_thai_formatting():
    document = Document()
    add_heading(document, 1, "บทที่ 4")
    heading = document.paragraphs[-1]
    properties = heading._p.get_or_add_pPr()

    assert heading.style.name == "Heading 1"
    assert properties.find(qn("w:outlineLvl")).get(qn("w:val")) == "0"
    assert properties.find(qn("w:pageBreakBefore")) is not None
    assert heading.runs[0].font.name == FONT_NAME
    assert heading.runs[0].font.size.pt == 20
    assert heading.runs[0].bold is True
    assert heading.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == FONT_NAME


def test_subheading_uses_heading_style_and_outline_without_forced_page_break():
    document = Document()
    add_heading(document, 2, "4.1 กรอบการประเมิน")
    heading = document.paragraphs[-1]
    properties = heading._p.get_or_add_pPr()

    assert heading.style.name == "Heading 2"
    assert properties.find(qn("w:outlineLvl")).get(qn("w:val")) == "1"
    assert properties.find(qn("w:pageBreakBefore")) is None
    assert heading.runs[0].font.name == FONT_NAME
    assert heading.runs[0].font.size.pt == 18


def test_injector_default_uses_dated_verified_chapter():
    assert DEFAULT_CHAPTER.name == "chapter_4_verified_results_20260807.docx"


def test_manifest_verification_rejects_tampered_markdown(tmp_path: Path):
    source = tmp_path / "chapter4.md"
    source.write_text("# บทที่ 4\n", encoding="utf-8")
    manifest_path = tmp_path / "manifest.json"
    manifest = {
        "status": "complete",
        "protocol": {
            "total_cases": 220,
            "train_cases": 125,
            "test_cases": 95,
            "standard_test_cases": 75,
            "adversarial_test_cases": 20,
        },
        "outputs": {
            "markdown": {
                "path": str(source),
                "sha256": sha256(source),
            }
        },
    }
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
    assert verify_source_manifest(source, manifest_path)["status"] == "complete"

    source.write_text("# altered\n", encoding="utf-8")
    with pytest.raises(ValueError, match="hash does not match"):
        verify_source_manifest(source, manifest_path)


def test_manifest_verification_rejects_missing_protocol_count(tmp_path: Path):
    source = tmp_path / "chapter4.md"
    source.write_text("# บทที่ 4\n", encoding="utf-8")
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(
        json.dumps({
            "status": "complete",
            "protocol": {
                "total_cases": 220,
                "train_cases": 125,
                "test_cases": 95,
                "standard_test_cases": 75,
            },
            "outputs": {
                "markdown": {"path": str(source), "sha256": sha256(source)}
            },
        }),
        encoding="utf-8",
    )
    with pytest.raises(ValueError, match="adversarial_test_cases"):
        verify_source_manifest(source, manifest_path)


def test_l2_stress_summary_reports_target_false_trigger_and_joint_rates():
    rows = [
        {
            "case_id": "ADV_A",
            "evaluation_slice": "adversarial_test",
            "expected_codes": ["TARGET_A"],
            "predicted_codes": ["TARGET_A"],
            "augmentation": {"false_trigger_code": "FALSE_A"},
            "detector_metrics": {"tp": 1, "fp": 0, "fn": 0, "f1": 1.0, "exact_match": 1.0},
            "l2_degraded": False,
        },
        {
            "case_id": "ADV_B",
            "evaluation_slice": "adversarial_test",
            "expected_codes": ["TARGET_B"],
            "predicted_codes": ["FALSE_B"],
            "augmentation": {"false_trigger_code": "FALSE_B"},
            "detector_metrics": {"tp": 0, "fp": 1, "fn": 1, "f1": 0.0, "exact_match": 0.0},
            "l2_degraded": False,
        },
    ]
    matrix = {"rows_by_model": {model: list(rows) for model in EXPECTED_MODELS}}
    summaries = l2_slice_summary(matrix, "adversarial_test")
    assert len(summaries) == 3
    for summary in summaries:
        assert summary["target_preservation_rate"] == 0.5
        assert summary["false_trigger_suppression_rate"] == 0.5
        assert summary["joint_pass_rate"] == 0.5
