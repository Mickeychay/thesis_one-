#!/usr/bin/env python3
"""Replace Chapter 4 in a thesis DOCX while preserving the surrounding package."""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CHAPTER = ROOT / "output" / "doc" / "chapter_4_verified_results_20260807.docx"


def normalized(value: str) -> str:
    return " ".join(value.split())


def element_text(element, document: Document) -> str:
    if element.tag == qn("w:p"):
        return normalized(Paragraph(element, document._body).text)
    return ""


def marker_elements(document: Document, marker: str) -> list:
    return [
        element
        for element in document.element.body
        if element.tag == qn("w:p") and element_text(element, document) == marker
    ]


def has_section_break(element) -> bool:
    return section_properties(element) is not None


def section_properties(element):
    if element.tag == qn("w:sectPr"):
        return element
    return element.find(f".//{qn('w:sectPr')}")


def ensure_heading_properties(element, level: int, *, page_break_before: bool = False) -> None:
    properties = element.get_or_add_pPr()
    style = properties.find(qn("w:pStyle"))
    if style is None:
        style = OxmlElement("w:pStyle")
        properties.insert(0, style)
    style.set(qn("w:val"), f"Heading{level}")

    outline = properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        properties.append(outline)
    outline.set(qn("w:val"), str(level - 1))

    if page_break_before:
        page_break = properties.find(qn("w:pageBreakBefore"))
        if page_break is None:
            page_break = OxmlElement("w:pageBreakBefore")
            properties.append(page_break)
        page_break.set(qn("w:val"), "true")


def clear_page_number_start(element) -> bool:
    properties = section_properties(element)
    if properties is None:
        return False
    page_number = properties.find(qn("w:pgNumType"))
    if page_number is None or page_number.get(qn("w:start")) is None:
        return False
    del page_number.attrib[qn("w:start")]
    return True


def chapter_elements(chapter: Document) -> list:
    if chapter.inline_shapes:
        raise ValueError(
            "The verified chapter contains inline shapes. Relationship-aware image "
            "copying is required before it can be injected safely."
        )
    elements = [
        element
        for element in chapter.element.body
        if element.tag != qn("w:sectPr")
    ]
    if not elements:
        raise ValueError("The verified chapter is empty")
    return elements


def validate_chapter(chapter: Document, start_marker: str, end_marker: str) -> None:
    starts = marker_elements(chapter, start_marker)
    ends = marker_elements(chapter, end_marker)
    if len(starts) != 1:
        raise ValueError(
            f"Verified chapter must contain exactly one {start_marker!r}; found {len(starts)}"
        )
    if ends:
        raise ValueError(f"Verified chapter must not contain the next marker {end_marker!r}")


def replace_chapter(
    source: Path,
    chapter_path: Path,
    output: Path,
    *,
    start_marker: str = "บทที่ 4",
    end_marker: str = "บทที่ 5",
) -> dict:
    source = source.resolve()
    chapter_path = chapter_path.resolve()
    output = output.resolve()
    document = Document(source)
    chapter = Document(chapter_path)
    validate_chapter(chapter, start_marker, end_marker)

    starts = marker_elements(document, start_marker)
    ends = marker_elements(document, end_marker)
    if len(starts) != 1 or len(ends) != 1:
        raise ValueError(
            f"Source must contain exactly one {start_marker!r} and one {end_marker!r}; "
            f"found {len(starts)} and {len(ends)}"
        )

    body = document.element.body
    body_elements = list(body)
    start_index = body_elements.index(starts[0])
    end_index = body_elements.index(ends[0])
    if start_index >= end_index:
        raise ValueError("Chapter markers are out of order")

    old_span = body_elements[start_index:end_index]
    preserved_section_breaks = [element for element in old_span if has_section_break(element)]
    insertion_anchor = preserved_section_breaks[0] if preserved_section_breaks else ends[0]
    removed_tables = sum(element.tag == qn("w:tbl") for element in old_span)
    source_sections = len(document.sections)
    source_tables = len(document.tables)

    chapter_start = marker_elements(chapter, start_marker)[0]
    ensure_heading_properties(chapter_start, 1, page_break_before=True)

    cleared_page_number_starts = sum(
        clear_page_number_start(element) for element in preserved_section_breaks
    )
    if preserved_section_breaks:
        for element in body_elements[end_index:]:
            if section_properties(element) is not None:
                cleared_page_number_starts += int(clear_page_number_start(element))
                break

    for element in old_span:
        if element not in preserved_section_breaks:
            body.remove(element)

    inserted = chapter_elements(chapter)
    for element in inserted:
        insertion_anchor.addprevious(deepcopy(element))

    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.stem}.tmp.docx")
    document.save(temporary)

    verified = Document(temporary)
    verified_starts = marker_elements(verified, start_marker)
    verified_ends = marker_elements(verified, end_marker)
    expected_tables = source_tables - removed_tables + len(chapter.tables)
    if len(verified_starts) != 1 or len(verified_ends) != 1:
        raise RuntimeError("Injected document failed chapter-marker validation")
    verified_body = list(verified.element.body)
    if verified_body.index(verified_starts[0]) >= verified_body.index(verified_ends[0]):
        raise RuntimeError("Injected Chapter 4 is not before Chapter 5")
    if len(verified.tables) != expected_tables:
        raise RuntimeError(
            f"Injected document has {len(verified.tables)} tables; expected {expected_tables}"
        )
    if len(verified.sections) != source_sections:
        raise RuntimeError(
            f"Injected document has {len(verified.sections)} sections; expected {source_sections}"
        )

    temporary.replace(output)
    return {
        "source": str(source),
        "chapter": str(chapter_path),
        "output": str(output),
        "removed_body_elements": len(old_span) - len(preserved_section_breaks),
        "preserved_section_breaks": len(preserved_section_breaks),
        "cleared_page_number_starts": cleared_page_number_starts,
        "inserted_body_elements": len(inserted),
        "tables": len(verified.tables),
        "sections": len(verified.sections),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--chapter", type=Path, default=DEFAULT_CHAPTER)
    parser.add_argument("--start-marker", default="บทที่ 4")
    parser.add_argument("--end-marker", default="บทที่ 5")
    args = parser.parse_args()
    result = replace_chapter(
        args.source,
        args.chapter,
        args.output,
        start_marker=args.start_marker,
        end_marker=args.end_marker,
    )
    for key, value in result.items():
        print(f"{key}: {value}")


if __name__ == "__main__":
    main()
