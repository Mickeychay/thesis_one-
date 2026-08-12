#!/usr/bin/env python3
"""Update the current thesis DOCX files with the 20-case adversarial slice."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from revise_thesis_chapters_2_3 import (
    ADVERSARIAL_METHOD_TEXT,
    DATASET_CURRENT_TEXT,
    SPLIT_CURRENT_TEXT,
    copy_paragraph_format,
    insert_before,
    normalized,
    set_paragraph_text,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATHS = [
    ROOT / "output" / "doc" / "65130641_Riskie_Thesis_v2_revised.docx",
    ROOT / "output" / "doc" / "65130641_Riskie_Thesis_v2_revised_final.docx",
]


def first_after(document: Document, heading_index: int, predicate):
    paragraphs = document.paragraphs
    for paragraph in paragraphs[heading_index + 1 : heading_index + 15]:
        if predicate(normalized(paragraph.text)):
            return paragraph
    raise RuntimeError("Could not locate the expected methodology paragraph")


def update_document(path: Path) -> None:
    document = Document(path)
    paragraphs = document.paragraphs
    heading_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if normalized(paragraph.text)
        == "3.9.1 การสร้างชุดข้อมูลอ้างอิง (Ground Truth Dataset)"
    )

    dataset_paragraph = first_after(
        document,
        heading_index,
        lambda text: "data/expanded_ground_truth.json" in text and "เคสต้นฉบับ" in text,
    )
    split_paragraph = first_after(
        document,
        heading_index,
        lambda text: "Family-Level Stratified Split" in text,
    )
    set_paragraph_text(dataset_paragraph, DATASET_CURRENT_TEXT)
    set_paragraph_text(split_paragraph, SPLIT_CURRENT_TEXT)

    adversarial_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if normalized(paragraph.text).startswith("เคสท้าทายระบบ (Adversarial Cases) คือ")
    ]
    if adversarial_paragraphs:
        copy_paragraph_format(dataset_paragraph, adversarial_paragraphs[0])
        set_paragraph_text(adversarial_paragraphs[0], ADVERSARIAL_METHOD_TEXT)
    else:
        paraphrase_paragraph = first_after(
            document,
            heading_index,
            lambda text: text.startswith("การถอดความ (Paraphrase) คือ"),
        )
        insert_before(
            paraphrase_paragraph,
            ADVERSARIAL_METHOD_TEXT,
            dataset_paragraph,
        )

    document.save(path)
    print(path)


def main() -> None:
    for path in DOCX_PATHS:
        if not path.exists():
            raise FileNotFoundError(path)
        update_document(path)


if __name__ == "__main__":
    main()
