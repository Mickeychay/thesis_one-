import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def replace_ch5_in_main_doc(main_path, ch5_path, out_path):
    main_doc = docx.Document(main_path)
    
    p_ch5_idx = -1
    for i, p in enumerate(main_doc.paragraphs):
        if 'บทที่ 5' in p.text:
            p_ch5_idx = i
            break
            
    print(f"Main doc paragraph index: Chapter 5 starts at {p_ch5_idx}")
    
    new_doc = docx.Document()
    
    # If Chapter 5 paragraph was found, copy up to Chapter 5 start
    limit = p_ch5_idx if p_ch5_idx != -1 else len(main_doc.paragraphs)
    for i in range(limit):
        p_old = main_doc.paragraphs[i]
        p_new = new_doc.add_paragraph()
        p_new.style = p_old.style
        p_new.alignment = p_old.alignment
        for run in p_old.runs:
            r = p_new.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            r.font.name = run.font.name
            r.font.size = run.font.size
            
    # Copy all elements (paragraphs and tables) from ch5_doc
    ch5_doc = docx.Document(ch5_path)
    for elem in ch5_doc.element.body:
        if elem.tag.endswith('p'):
            p_obj = docx.text.paragraph.Paragraph(elem, ch5_doc)
            p_new = new_doc.add_paragraph()
            p_new.alignment = p_obj.alignment
            p_new.paragraph_format.space_before = p_obj.paragraph_format.space_before
            p_new.paragraph_format.space_after = p_obj.paragraph_format.space_after
            for run in p_obj.runs:
                r = p_new.add_run(run.text)
                r.bold = run.bold
                r.italic = run.italic
                r.font.name = 'TH Sarabun PSK'
                r.font.size = run.font.size if run.font.size else Pt(16)
                if run.font.color and run.font.color.rgb:
                    r.font.color.rgb = run.font.color.rgb
        elif elem.tag.endswith('tbl'):
            tbl_obj = docx.table.Table(elem, ch5_doc)
            new_table = new_doc.add_table(rows=len(tbl_obj.rows), cols=len(tbl_obj.columns))
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row in enumerate(tbl_obj.rows):
                for c_idx, cell in enumerate(row.cells):
                    target_cell = new_table.rows[r_idx].cells[c_idx]
                    target_cell.text = cell.text
                    set_cell_margins(target_cell, top=100, bottom=100, left=150, right=150)
                    if r_idx == 0:
                        set_cell_background(target_cell, "F0F4F8")
                    else:
                        bg = "FAFAFA" if r_idx % 2 == 1 else "FFFFFF"
                        set_cell_background(target_cell, bg)
                    p = target_cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = 'TH Sarabun PSK'
                        r.font.size = Pt(14)
                        if r_idx == 0:
                            r.font.bold = True

    new_doc.save(out_path)
    print(f"Successfully injected Chapter 5 into main docx: {out_path}")

if __name__ == '__main__':
    main_thesis = ROOT / '65130641_Riskie_Thesis_Glen_fixed.docx'
    ch5_docx = ROOT / 'ch5_thesis.docx'
    if main_thesis.is_file():
        replace_ch5_in_main_doc(main_thesis, ch5_docx, main_thesis)
