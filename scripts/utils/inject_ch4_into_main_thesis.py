"""Replace Chapter 4 of the main thesis DOCX with the freshly built ch4_thesis.docx.

Implementation note — why this edits the document in place:

The previous version built a blank `docx.Document()` and re-created every
paragraph by copying `run.text` plus a few font attributes. That silently
discarded everything not expressed as run text:

  * embedded images (a run's picture lives in the run XML, not in `run.text`),
    so all four Chapter 4 figures vanished on every injection;
  * tables outside Chapter 4, because `main_doc.paragraphs` skips tables —
    this dropped the Chapter 5 evidence-summary table;
  * section properties, headers/footers, numbering, and the original styles.

Editing the existing body avoids all of that: only the Chapter 4 element range
is removed, and the new Chapter 4 elements are spliced into the same position.
Everything outside that range keeps its original XML byte-for-byte.
"""

import copy

import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn


CH4_HEADING = 'บทที่ 4'
CH5_HEADING = 'บทที่ 5'


def _body_position(body_children, paragraph):
    """Index of a Paragraph's XML element within the body's children."""
    return body_children.index(paragraph._p)


def _find_chapter_bounds(doc):
    """Return (ch4_element, ch5_element) for the chapter heading paragraphs."""
    ch4 = ch5 = None
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if CH4_HEADING in text and ch4 is None:
            ch4 = paragraph
        elif CH5_HEADING in text and ch5 is None:
            ch5 = paragraph
            break
    if ch4 is None:
        raise ValueError(f"could not locate {CH4_HEADING!r} in the main document")
    if ch5 is None:
        raise ValueError(f"could not locate {CH5_HEADING!r} in the main document")
    return ch4, ch5


def _remap_images(element, source_part, target_part):
    """Re-point r:embed / r:link ids from source_part's rels to target_part's.

    A deep-copied run keeps the source document's relationship id. Left
    unremapped, Word reports the file as corrupt or shows an empty frame, so
    every image part is added to the target package and the id rewritten.
    """
    embed_attr = qn('r:embed')
    link_attr = qn('r:link')

    for blip in element.iter(qn('a:blip')):
        for attr in (embed_attr, link_attr):
            rid = blip.get(attr)
            if not rid:
                continue
            source_rel = source_part.rels[rid]
            if source_rel.is_external:
                new_rid = target_part.relate_to(
                    source_rel.target_ref, source_rel.reltype, is_external=True
                )
            else:
                image_part = source_rel.target_part
                new_rid = target_part.relate_to(image_part, RT.IMAGE)
            blip.set(attr, new_rid)


def replace_ch4_in_main_doc(main_path, out_path, ch4_path='ch4_thesis.docx'):
    main_doc = docx.Document(main_path)
    ch4_doc = docx.Document(ch4_path)

    ch4_heading, ch5_heading = _find_chapter_bounds(main_doc)

    body = main_doc.element.body
    children = list(body)
    start = _body_position(children, ch4_heading)
    end = _body_position(children, ch5_heading)
    if start >= end:
        raise ValueError(
            f"{CH4_HEADING!r} (pos {start}) must precede {CH5_HEADING!r} (pos {end})"
        )

    old_elements = children[start:end]
    print(f"Replacing {len(old_elements)} body elements "
          f"(positions {start}–{end - 1}) with new Chapter 4")

    # Splice new Chapter 4 content in before the Chapter 5 heading, preserving
    # document order, then drop the old range.
    anchor = ch5_heading._p
    copied_images = 0
    inserted = 0
    for element in ch4_doc.element.body:
        # Skip the trailing <w:sectPr>: section setup belongs to the main
        # document and must not be overwritten by the fragment's defaults.
        if element.tag == qn('w:sectPr'):
            continue
        new_element = copy.deepcopy(element)
        before = len(main_doc.part.rels)
        _remap_images(new_element, ch4_doc.part, main_doc.part)
        copied_images += len(main_doc.part.rels) - before
        anchor.addprevious(new_element)
        inserted += 1

    for element in old_elements:
        body.remove(element)

    print(f"Inserted {inserted} elements, remapped {copied_images} image relationships")

    main_doc.save(out_path)
    print(f"Successfully injected Chapter 4 into main docx: {out_path}")


if __name__ == '__main__':
    replace_ch4_in_main_doc(
        '65130641_Riskie_Thesis_Glen_fixed.docx',
        '65130641_Riskie_Thesis_Glen_fixed.docx',
    )
