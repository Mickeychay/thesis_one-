import re
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "md_report" / "thesis_full_ch4_final.md"
OUTPUT_DOCX = ROOT / "ch4_thesis.docx"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text, style='Normal', space_before=0, space_after=4, line_spacing=1.15, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, color_rgb=(0,0,0), font_size=16):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        run = p.add_run()
        run.font.name = 'TH Sarabun PSK'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*color_rgb)
        
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(font_size - 1)
            run.font.color.rgb = RGBColor(180, 40, 40)
        else:
            run.text = token
            run.bold = bold
            run.italic = italic
            
    return p

def convert_md_to_docx(md_path, out_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Markdown table handling
        if '|' in line and not in_code_block:
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            in_table = False
            # Render Markdown table
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            data_rows = []
            for row_line in table_lines[2:]: # Skip separator row
                if '|' in row_line:
                    row_cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    if row_cells:
                        data_rows.append(row_cells)
                        
            if headers and data_rows:
                table = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # Header row
                for c_idx, h_text in enumerate(headers):
                    cell = table.rows[0].cells[c_idx]
                    cell.text = h_text
                    set_cell_background(cell, "F0F4F8")
                    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = 'TH Sarabun PSK'
                        r.font.size = Pt(14)
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0, 51, 102)
                # Data rows
                for r_idx, r_cells in enumerate(data_rows):
                    for c_idx, cell_value in enumerate(r_cells):
                        if c_idx < len(headers):
                            cell = table.rows[r_idx+1].cells[c_idx]
                            cell.text = cell_value
                            bg = "FAFAFA" if r_idx % 2 == 1 else "FFFFFF"
                            set_cell_background(cell, bg)
                            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.font.name = 'TH Sarabun PSK'
                                r.font.size = Pt(13)
            table_lines = []
            
        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "F4F6F8")
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)
                crun = cp.add_run(code_text)
                crun.font.name = 'Courier New'
                crun.font.size = Pt(11)
                crun.font.color.rgb = RGBColor(40, 40, 40)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
            
        if not line.strip():
            i += 1
            continue
            
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            p._element.get_or_add_pPr().append(pBdr)
            i += 1
            continue
            
        if line.startswith('# '):
            add_styled_paragraph(doc, line[2:], space_before=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color_rgb=(0, 51, 102), font_size=22)
        elif line.startswith('## '):
            add_styled_paragraph(doc, line[3:], space_before=10, space_after=4, bold=True, color_rgb=(0, 51, 102), font_size=18)
        elif line.startswith('### '):
            add_styled_paragraph(doc, line[4:], space_before=8, space_after=3, bold=True, color_rgb=(0, 51, 102), font_size=16)
        elif line.startswith('#### '):
            add_styled_paragraph(doc, line[5:], space_before=6, space_after=2, bold=True, color_rgb=(51, 51, 51), font_size=16)
        elif line.startswith('> '):
            quote_text = line[2:]
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            set_cell_background(cell, "F0F4F8")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            tcPr = cell._element.get_or_add_tcPr()
            borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="003366"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
            tcPr.append(borders)
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            crun = cp.add_run(quote_text)
            crun.font.name = 'TH Sarabun PSK'
            crun.font.size = Pt(15)
            crun.font.italic = True
            crun.font.color.rgb = RGBColor(30, 30, 30)
        elif line.startswith('* ') or line.startswith('- '):
            add_styled_paragraph(doc, "• " + line[2:], space_before=1, space_after=3, font_size=16)
        elif re.match(r'^\d+\.\s', line):
            add_styled_paragraph(doc, line, space_before=1, space_after=3, font_size=16)
        else:
            add_styled_paragraph(doc, line, space_before=0, space_after=4, font_size=16)
            
        i += 1

    # Check if trailing table needs rendering
    if in_table and table_lines:
        headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
        data_rows = []
        for row_line in table_lines[2:]:
            if '|' in row_line:
                row_cells = [c.strip() for c in row_line.split('|')[1:-1]]
                if row_cells:
                    data_rows.append(row_cells)
        if headers and data_rows:
            table = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c_idx, h_text in enumerate(headers):
                cell = table.rows[0].cells[c_idx]
                cell.text = h_text
                set_cell_background(cell, "F0F4F8")
                set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'TH Sarabun PSK'
                    r.font.size = Pt(14)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0, 51, 102)
            for r_idx, r_cells in enumerate(data_rows):
                for c_idx, cell_value in enumerate(r_cells):
                    if c_idx < len(headers):
                        cell = table.rows[r_idx+1].cells[c_idx]
                        cell.text = cell_value
                        bg = "FAFAFA" if r_idx % 2 == 1 else "FFFFFF"
                        set_cell_background(cell, bg)
                        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.font.name = 'TH Sarabun PSK'
                            r.font.size = Pt(13)
        
    doc.save(out_path)
    print(f"Successfully generated DOCX at {out_path}")

if __name__ == '__main__':
    convert_md_to_docx(MD_PATH, OUTPUT_DOCX)
