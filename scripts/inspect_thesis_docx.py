from __future__ import annotations

import json
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def twips_to_cm(value):
    if value is None:
        return None
    return round(value.cm, 3)


def paragraph_summary(paragraph, index):
    p_pr = paragraph._p.pPr
    page_break_before = False
    keep_with_next = False
    if p_pr is not None:
        page_break_before = p_pr.find(qn("w:pageBreakBefore")) is not None
        keep_with_next = p_pr.find(qn("w:keepNext")) is not None

    return {
        "index": index,
        "text": paragraph.text[:240],
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment),
        "page_break_before": page_break_before,
        "keep_with_next": keep_with_next,
        "runs": [
            {
                "text": run.text[:80],
                "font": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
                "bold": run.bold,
                "italic": run.italic,
            }
            for run in paragraph.runs[:8]
        ],
    }


def inspect(path: Path):
    doc = Document(path)
    paragraphs = doc.paragraphs
    styles = Counter(
        paragraph.style.name if paragraph.style else "(none)"
        for paragraph in paragraphs
    )
    sections = []
    for i, section in enumerate(doc.sections):
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(qn("w:pgNumType"))
        sections.append(
            {
                "index": i,
                "start_type": str(section.start_type),
                "page_width_cm": twips_to_cm(section.page_width),
                "page_height_cm": twips_to_cm(section.page_height),
                "top_margin_cm": twips_to_cm(section.top_margin),
                "bottom_margin_cm": twips_to_cm(section.bottom_margin),
                "left_margin_cm": twips_to_cm(section.left_margin),
                "right_margin_cm": twips_to_cm(section.right_margin),
                "header_distance_cm": twips_to_cm(section.header_distance),
                "footer_distance_cm": twips_to_cm(section.footer_distance),
                "different_first_page": section.different_first_page_header_footer,
                "header_linked": section.header.is_linked_to_previous,
                "footer_linked": section.footer.is_linked_to_previous,
                "page_number_start": (
                    pg_num_type.get(qn("w:start")) if pg_num_type is not None else None
                ),
                "page_number_format": (
                    pg_num_type.get(qn("w:fmt")) if pg_num_type is not None else None
                ),
                "header_text": " | ".join(p.text for p in section.header.paragraphs),
                "footer_text": " | ".join(p.text for p in section.footer.paragraphs),
            }
        )

    heading_markers = []
    needles = ("บทที่ 1", "บทที่ 2", "บทที่ 3", "บทที่ 4", "บทที่ 5")
    for i, paragraph in enumerate(paragraphs):
        normalized = " ".join(paragraph.text.split())
        if any(needle in normalized for needle in needles):
            heading_markers.append(paragraph_summary(paragraph, i))

    payload = {
        "path": str(path),
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "inline_shape_count": len(doc.inline_shapes),
        "styles": styles.most_common(),
        "sections": sections,
        "heading_markers": heading_markers,
        "tail": [
            paragraph_summary(paragraph, i)
            for i, paragraph in list(enumerate(paragraphs))[-24:]
        ],
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    inspect(Path(sys.argv[1]))
