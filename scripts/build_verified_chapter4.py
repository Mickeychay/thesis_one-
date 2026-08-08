#!/usr/bin/env python3
"""Build the verified Chapter 4 DOCX from a hash-locked Markdown source."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "md_report" / "thesis_ch4_verified_20260807.md"
MANIFEST = ROOT / "evaluation_results" / "chapter4_artifact_manifest_20260807.json"
OUTPUT = ROOT / "output" / "doc" / "chapter_4_verified_results_20260807.docx"
FONT_NAME = "Sarabun"


def set_thai_run(run, *, size: float, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:eastAsia"), "th-TH")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, value: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_begin, instr, fld_end))
    set_thai_run(run, size=12)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str, size: float, *, default_bold: bool = False) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        bold = default_bold
        italic = False
        value = part
        if part.startswith("**") and part.endswith("**"):
            value = part[2:-2]
            bold = True
        elif part.startswith("`") and part.endswith("`"):
            value = part[1:-1]
            italic = True
        run = paragraph.add_run(value)
        set_thai_run(run, size=size, bold=bold, italic=italic)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(1.25)
    fmt.line_spacing = 1.0
    fmt.space_after = Pt(4)
    fmt.widow_control = True
    add_inline(paragraph, text, 16)


def set_outline_level(paragraph, level: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    outline = properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        properties.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def add_heading(doc: Document, level: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles[f"Heading {level}"]
    set_outline_level(paragraph, level)
    fmt = paragraph.paragraph_format
    fmt.keep_with_next = True
    fmt.widow_control = True
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if " ".join(text.split()) == "บทที่ 4":
            fmt.page_break_before = True
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(2)
        size = 20
    elif level == 2 and not text.startswith("4."):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(12)
        size = 20
    elif level == 2:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(4)
        size = 18
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(8)
        fmt.space_after = Pt(3)
        size = 16
    add_inline(paragraph, text, size, default_bold=True)


def split_markdown_row(line: str) -> list[str]:
    """Split a Markdown row without treating escaped pipes as delimiters."""
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for character in value:
        if escaped:
            current.append(character)
            escaped = False
        elif character == "\\":
            escaped = True
        elif character == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(character)
    if escaped:
        current.append("\\")
    cells.append("".join(current).strip())
    return cells


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        rows.append(split_markdown_row(line))
    if len(rows) < 2:
        raise ValueError("Markdown table requires a header and separator")
    return rows[0], rows[2:]


def clean_table_value(value: str) -> tuple[str, bool, bool]:
    bold = value.startswith("**") and value.endswith("**")
    italic = value.startswith("`") and value.endswith("`")
    if bold:
        value = value[2:-2]
    elif italic:
        value = value[1:-1]
    return value, bold, italic


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    font_size = 9.5 if len(headers) >= 9 else 11 if len(headers) >= 7 else 12

    for col, value in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9E2F3")
        set_cell_margins(cell)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        add_inline(paragraph, value, font_size, default_bold=True)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])

    numeric = re.compile(r"^[+\-]?(?:\d|\.)")
    for row_index, values in enumerate(rows, start=1):
        row = table.add_row()
        prevent_row_split(row)
        for col in range(len(headers)):
            raw = values[col] if col < len(values) else ""
            value, bold, italic = clean_table_value(raw)
            cell = row.cells[col]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index % 2 == 0:
                set_cell_shading(cell, "F5F7FA")
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if col > 0 and numeric.match(value) else
                WD_ALIGN_PARAGRAPH.LEFT if col == 0 else
                WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(value)
            set_thai_run(run, size=font_size, bold=bold, italic=italic)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)
    set_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(16)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def display_path(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(ROOT))
    except ValueError:
        return str(path.resolve())


def resolve_manifest_path(value: str) -> Path:
    path = Path(value)
    return path.resolve() if path.is_absolute() else (ROOT / path).resolve()


def verify_source_manifest(source: Path, manifest_path: Path) -> dict:
    if not source.is_file():
        raise FileNotFoundError(f"Verified Chapter 4 source is missing: {source}")
    if not manifest_path.is_file():
        raise FileNotFoundError(f"Chapter 4 evidence manifest is missing: {manifest_path}")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if manifest.get("status") != "complete":
        raise ValueError("Chapter 4 evidence manifest status must be complete")
    protocol = manifest.get("protocol", {})
    expected_counts = {
        "total_cases": 220,
        "train_cases": 125,
        "test_cases": 95,
        "standard_test_cases": 75,
        "adversarial_test_cases": 20,
    }
    for field, expected in expected_counts.items():
        if int(protocol.get(field, -1)) != expected:
            raise ValueError(f"Manifest protocol {field} does not equal {expected}")
    markdown = manifest.get("outputs", {}).get("markdown", {})
    reported_path = markdown.get("path")
    reported_hash = markdown.get("sha256")
    if not reported_path or resolve_manifest_path(str(reported_path)) != source.resolve():
        raise ValueError("Manifest Markdown path does not match the requested source")
    actual_hash = sha256(source)
    if reported_hash != actual_hash:
        raise ValueError(
            "Manifest Markdown hash does not match the requested source: "
            f"manifest={reported_hash!r}, actual={actual_hash}"
        )
    return manifest


def update_docx_manifest(manifest: dict, manifest_path: Path, output: Path) -> None:
    manifest["docx_generated_at"] = datetime.now(ZoneInfo("Asia/Bangkok")).isoformat(timespec="seconds")
    manifest["docx_builder"] = {
        "path": display_path(Path(__file__)),
        "sha256": sha256(Path(__file__)),
    }
    manifest.setdefault("outputs", {})["docx"] = {
        "path": display_path(output),
        "sha256": sha256(output),
        "bytes": output.stat().st_size,
        "source_markdown_sha256": manifest["outputs"]["markdown"]["sha256"],
    }
    temporary = manifest_path.with_suffix(f"{manifest_path.suffix}.tmp")
    temporary.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    os.replace(temporary, manifest_path)


def build(source: Path = SOURCE, output: Path = OUTPUT, manifest_path: Path = MANIFEST) -> None:
    source = source.resolve()
    output = output.resolve()
    manifest_path = manifest_path.resolve()
    manifest = verify_source_manifest(source, manifest_path)
    doc = Document()
    configure_document(doc)
    lines = source.read_text(encoding="utf-8").splitlines()
    pending_table_title: str | None = None
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("### "):
            add_heading(doc, 3, line[4:])
            index += 1
            continue
        if line.startswith("## "):
            add_heading(doc, 2, line[3:])
            index += 1
            continue
        if line.startswith("# "):
            add_heading(doc, 1, line[2:])
            index += 1
            continue
        if line.startswith("**ตารางที่") and line.endswith("**"):
            pending_table_title = line[2:-2]
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            if pending_table_title:
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.paragraph_format.keep_with_next = True
                title.paragraph_format.space_before = Pt(6)
                title.paragraph_format.space_after = Pt(3)
                add_inline(title, pending_table_title, 16, default_bold=True)
                pending_table_title = None
            headers, rows = parse_table(table_lines)
            add_table(doc, headers, rows)
            continue
        add_body_paragraph(doc, line)
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    update_docx_manifest(manifest, manifest_path, output)
    print(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--manifest", type=Path, default=MANIFEST)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    build(args.source, args.output, args.manifest)
